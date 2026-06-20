from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ============ 样式设置 ============
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# ============ 标题 ============
title = doc.add_heading('RT-DETR 项目文件结构详解', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('')

# ============ 1. 项目总览 ============
doc.add_heading('一、项目总览', level=1)
p = doc.add_paragraph()
p.add_run('这是基于 ').font.size = Pt(11)
p.add_run('Ultralytics YOLOv8').bold = True
p.add_run(' 框架改造的 ').font.size = Pt(11)
p.add_run('RT-DETR (Real-Time DEtection TRansformer)').bold = True
p.add_run(' 目标检测项目，版本号 ').font.size = Pt(11)
p.add_run('8.0.201').bold = True
p.add_run('。它在 Ultralytics 框架基础上整合了大量创新模块（注意力机制、新型卷积、Mamba 架构等），供论文研究和实验使用。').font.size = Pt(11)

# ============ 2. 根目录入口脚本 ============
doc.add_heading('二、根目录入口脚本', level=1)

table1 = doc.add_table(rows=6, cols=2, style='Light Grid Accent 1')
table1.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table1.rows[0].cells
hdr[0].text = '文件'
hdr[1].text = '作用'

data1 = [
    ('train.py', '训练入口。加载 RT-DETR 模型的 YAML 配置文件，调用 model.train() 启动训练。支持单卡/多卡训练、断点恢复 (resume)、早停 (patience) 等参数设置。'),
    ('detect.py', '推理/检测入口。加载训练好的 .pt 权重对图片/视频进行目标检测，输出带有检测框的图片。支持置信度阈值、线宽、标签显示等配置。'),
    ('val.py', '验证/评估入口。加载训练好的权重，在验证集/测试集上计算精度(Precision)、召回率(Recall)、mAP50、mAP75、mAP50-95、F1-Score、参数量、GFLOPs、FPS 等指标，并生成论文格式的数据表格保存到 paper_data.txt。'),
    ('export.py', '模型导出入口。将训练好的 .pt 模型导出为 ONNX 或 TensorRT (engine) 格式，便于推理部署。'),
    ('track.py', '目标跟踪入口。使用 RT-DETR 模型对视频进行多目标跟踪 (MOT)，基于 Bot-SORT / ByteTrack 算法。'),
]
for i, (file, desc) in enumerate(data1):
    row = table1.rows[i + 1]
    row.cells[0].text = file
    row.cells[1].text = desc

doc.add_paragraph('')

# ============ 3. 辅助工具脚本 ============
doc.add_heading('三、辅助工具脚本', level=1)

table2 = doc.add_table(rows=9, cols=2, style='Light Grid Accent 1')
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr2 = table2.rows[0].cells
hdr2[0].text = '文件'
hdr2[1].text = '作用'

data2 = [
    ('get_COCO_metrice.py', 'COCO 指标计算。使用 pycocotools 和 TIDE 工具，输入标注 JSON 和预测 JSON，计算 COCO 标准的 mAP 等指标。'),
    ('get_FPS.py', 'FPS/推理速度测试。对模型进行预热后，在指定 batch size 下测试推理延迟和 FPS，支持 FP16 模式。'),
    ('get_all_yaml_param_and_flops.py', '批量 YAML 参数量统计。遍历 ultralytics/cfg/models/rt-detr/ 下所有 YAML 配置文件，逐一计算每个模型的 GFLOPs 和参数量并排序输出。'),
    ('get_model_erf.py', '有效感受野 (ERF) 可视化。通过计算输出对输入梯度的贡献分数，生成模型有效感受野的热力图，评估模型关注区域的大小。'),
    ('heatmap.py', '热力图可视化。使用 Grad-CAM 系列方法 (GradCAM++/XGradCAM/EigenCAM 等) 对 RT-DETR 模型的中间层生成特征热力图，叠加检测框。'),
    ('main_profile.py', '模型结构分析。加载 YAML 配置的模型，输出每层的详细信息（输入输出形状、参数量、计算量）和模型整体信息。'),
    ('plot_result.py', '训练曲线绘制。读取训练产生的 results.csv 文件，绘制精度/召回率/mAP 曲线和训练/验证 loss 曲线，保存为 PNG 图片。'),
    ('test_env.py', '环境测试。验证各种依赖模块是否正确安装（mmcv、mamba、DCNv3/v4、KAT、natten 等）。'),
]
for i, (file, desc) in enumerate(data2):
    row = table2.rows[i + 1]
    row.cells[0].text = file
    row.cells[1].text = desc

doc.add_paragraph('')

# ============ 4. 配置文件 ============
doc.add_heading('四、配置文件', level=1)

table3 = doc.add_table(rows=4, cols=2, style='Light Grid Accent 1')
table3.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr3 = table3.rows[0].cells
hdr3[0].text = '文件'
hdr3[1].text = '作用'

data3 = [
    ('requirements.txt', 'Python 依赖列表：PyTorch、OpenCV、NumPy、Matplotlib 等基础依赖。'),
    ('setup.py', '将 ultralytics 安装为 Python 包，支持 pip install，注册了 yolo 和 ultralytics 命令行入口。'),
    ('setup.cfg', 'setuptools 元数据配置。'),
]
for i, (file, desc) in enumerate(data3):
    row = table3.rows[i + 1]
    row.cells[0].text = file
    row.cells[1].text = desc

doc.add_paragraph('')

# ============ 5. 核心库 ============
doc.add_heading('五、核心库 ultralytics/', level=1)

# 5.1
doc.add_heading('5.1 __init__.py', level=2)
doc.add_paragraph('包的入口，导出版本号 8.0.201 和五类模型：YOLO、NAS、SAM、FastSAM、RTDETR。')

# 5.2
doc.add_heading('5.2 engine/ — 引擎层（核心流程）', level=2)

table4 = doc.add_table(rows=8, cols=2, style='Light Grid Accent 1')
table4.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr4 = table4.rows[0].cells
hdr4[0].text = '文件'
hdr4[1].text = '作用'

data4 = [
    ('model.py', '所有模型的基类，提供 train()、val()、predict()、export()、track() 等统一接口。'),
    ('trainer.py', '训练引擎，管理训练循环、优化器、学习率调度、混合精度训练、EMA、DDP 分布式训练等。'),
    ('validator.py', '验证引擎，在验证集上计算各类指标（mAP、Precision、Recall 等）。'),
    ('predictor.py', '推理引擎，处理图片/视频的预处理、前向推理、后处理、结果保存。'),
    ('exporter.py', '模型导出引擎，支持导出 ONNX、TensorRT、OpenVINO、CoreML、TFLite 等格式。'),
    ('results.py', '结果类，封装检测/分割/跟踪的结果（boxes、masks、keypoints 等）。'),
    ('tuner.py', '超参数调优，使用 Ray Tune 自动搜索最佳超参数。'),
]
for i, (file, desc) in enumerate(data4):
    row = table4.rows[i + 1]
    row.cells[0].text = file
    row.cells[1].text = desc

doc.add_paragraph('')

# 5.3
doc.add_heading('5.3 nn/ — 神经网络层', level=2)

table5 = doc.add_table(rows=7, cols=2, style='Light Grid Accent 1')
table5.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr5 = table5.rows[0].cells
hdr5[0].text = '文件/目录'
hdr5[1].text = '作用'

data5 = [
    ('tasks.py', '核心任务模块，解析 YAML 配置文件构建模型结构，实现 attempt_load_weights 加载权重。'),
    ('autobackend.py', '自动后端，支持多种推理后端 (PyTorch、ONNX、TensorRT、OpenVINO 等)。'),
    ('modules/', '基础模块：block.py (C2f/C3/Bottleneck 等)、conv.py (各种卷积)、head.py (检测头)、transformer.py (Transformer 编解码器)。'),
    ('backbone/', '18 种骨干网络：SwinTransformer、ConvNeXtV2、EfficientFormerV2、VanillaNet、RepViT、StarNet、MambaOut、TransNext、UniRepLKNet、LSKNet、PKINet、MobileNetV4、FasterNet 等。'),
    ('extra_modules/', '80+ 个创新模块：涵盖注意力机制(CBAM/SE/ECA/SimAM/EMA等)、特征融合(ASFF/BiFPN/SDI等)、新型卷积(DCNv3/v4/DynamicSnakeConv/WTConv等)、Transformer/Mamba结构(MambaVision/TransMamba/GradMamba等)、轻量化设计(GhostConv/RepConv/PartialConv等)、频域模块(FcaNet/FreqFusion/WaveletPool等)。'),
    ('cfg/', 'YAML 配置文件目录，包含模型结构定义（RT-DETR 各变体的 YAML）和训练数据配置。'),
]
for i, (file, desc) in enumerate(data5):
    row = table5.rows[i + 1]
    row.cells[0].text = file
    row.cells[1].text = desc

doc.add_paragraph('')

# 5.4
doc.add_heading('5.4 models/ — 各模型实现', level=2)

table6 = doc.add_table(rows=7, cols=2, style='Light Grid Accent 1')
table6.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr6 = table6.rows[0].cells
hdr6[0].text = '文件/目录'
hdr6[1].text = '作用'

data6 = [
    ('rtdetr/', 'RT-DETR 模型：model.py（模型定义与训练逻辑）、predict.py（预测逻辑）、train.py（训练专用逻辑）、val.py（验证逻辑）。'),
    ('yolo/', 'YOLO 系列模型实现。'),
    ('sam/', 'SAM (Segment Anything Model) 分割模型。'),
    ('fastsam/', 'FastSAM 轻量分割模型。'),
    ('nas/', 'NAS 神经架构搜索模型。'),
    ('utils/', '损失函数 (VarifocalLoss、FocalLoss、BboxLoss、DetrLoss) 和匈牙利匹配器 (HungarianMatcher) 等 DETR 特有操作。'),
]
for i, (file, desc) in enumerate(data6):
    row = table6.rows[i + 1]
    row.cells[0].text = file
    row.cells[1].text = desc

doc.add_paragraph('')

# 5.5
doc.add_heading('5.5 data/ — 数据处理', level=2)

table7 = doc.add_table(rows=7, cols=2, style='Light Grid Accent 1')
table7.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr7 = table7.rows[0].cells
hdr7[0].text = '文件'
hdr7[1].text = '作用'

data7 = [
    ('dataset.py', '自定义数据集类，加载图片和标签。'),
    ('augment.py', '数据增强（Mosaic、MixUp、HSV 变换、翻转、缩放等）。'),
    ('base.py', '数据加载基类。'),
    ('build.py', '构建 DataLoader，分布式采样。'),
    ('loaders.py', '多种数据源加载器（图片/视频/流媒体）。'),
    ('converter.py', '数据集格式转换工具。'),
]
for i, (file, desc) in enumerate(data7):
    row = table7.rows[i + 1]
    row.cells[0].text = file
    row.cells[1].text = desc

doc.add_paragraph('')

# 5.6
doc.add_heading('5.6 utils/ — 工具函数', level=2)

table8 = doc.add_table(rows=11, cols=2, style='Light Grid Accent 1')
table8.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr8 = table8.rows[0].cells
hdr8[0].text = '文件'
hdr8[1].text = '作用'

data8 = [
    ('loss.py', '训练损失函数。'),
    ('metrics.py', '目标检测评估指标计算（mAP、IoU、混淆矩阵等）。'),
    ('ops.py', '通用操作（NMS、坐标变换 xywh↔xyxy、缩放框等）。'),
    ('tal.py', 'Task-Aligned Assigner 标签分配策略。'),
    ('plotting.py', '结果可视化（画框、标签、颜色映射）。'),
    ('torch_utils.py', 'PyTorch 工具（设备选择、模型参数统计 model_info、模型融合、FP16 转换等）。'),
    ('checks.py', '环境检查（依赖版本、GPU 可用性、数据完整性）。'),
    ('downloads.py', '从网络下载预训练权重和配置文件。'),
    ('callbacks/', '回调函数，支持 TensorBoard、ClearML、Comet、MLflow、Neptune、W&B 等训练日志记录平台。'),
    ('tuner.py', '超参数搜索辅助。'),
]
for i, (file, desc) in enumerate(data8):
    row = table8.rows[i + 1]
    row.cells[0].text = file
    row.cells[1].text = desc

doc.add_paragraph('')

# 5.7
doc.add_heading('5.7 其他核心目录', level=2)

table9 = doc.add_table(rows=4, cols=2, style='Light Grid Accent 1')
table9.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr9 = table9.rows[0].cells
hdr9[0].text = '目录'
hdr9[1].text = '作用'

data9 = [
    ('cfg/', 'YAML 配置文件目录，包含模型结构定义（RT-DETR 各变体的 YAML）和训练数据配置。'),
    ('hub/', 'Ultralytics HUB 集成（云端训练/部署），含认证(auth.py)和会话管理(session.py)。'),
    ('trackers/', '多目标跟踪算法实现：ByteTrack (byte_tracker.py)、Bot-SORT (bot_sort.py)、卡尔曼滤波、匈牙利匹配等。'),
]
for i, (file, desc) in enumerate(data9):
    row = table9.rows[i + 1]
    row.cells[0].text = file
    row.cells[1].text = desc

doc.add_paragraph('')

# ============ 6. dataset目录 ============
doc.add_heading('六、dataset/ 目录', level=1)

table10 = doc.add_table(rows=4, cols=2, style='Light Grid Accent 1')
table10.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr10 = table10.rows[0].cells
hdr10[0].text = '目录'
hdr10[1].text = '作用'

data10 = [
    ('dataset/images/', '训练/验证/测试图片存放目录。'),
    ('dataset/labels/', '对应的 YOLO 格式标注文件存放目录。'),
    ('dataset/VOCdevkit/', 'VOC 格式数据集（含 Annotations、JPEGImages、txt 标签）。'),
]
for i, (file, desc) in enumerate(data10):
    row = table10.rows[i + 1]
    row.cells[0].text = file
    row.cells[1].text = desc

doc.add_paragraph('')

# ============ 7. 文档类文件 ============
doc.add_heading('七、文档类文件', level=1)

table11 = doc.add_table(rows=6, cols=2, style='Light Grid Accent 1')
table11.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr11 = table11.rows[0].cells
hdr11[0].text = '文件'
hdr11[1].text = '作用'

data11 = [
    ('README.md', '项目说明文档。'),
    ('使用教程.md', '中文使用教程，含常见错误和解决方案。'),
    ('LOSS改进系列.md', '损失函数改进的系列文档。'),
    ('论文辅导项目介绍.md', '论文辅导项目说明。'),
    ('项目百度云视频.txt', '配套视频教程的百度云链接。'),
]
for i, (file, desc) in enumerate(data11):
    row = table11.rows[i + 1]
    row.cells[0].text = file
    row.cells[1].text = desc

doc.add_paragraph('')

# ============ 8. 总结 ============
doc.add_heading('八、总结：数据流和工作流', level=1)

doc.add_paragraph(
    'YAML配置 → tasks.py 解析 → 构建模型\n'
    '    ↓\n'
    'train.py → trainer.py 训练 → 保存 best.pt\n'
    '    ↓\n'
    'val.py → validator.py 验证 → 论文数据表格\n'
    '    ↓\n'
    'detect.py → predictor.py 推理 → 检测结果图片\n'
    '    ↓\n'
    'export.py → ONNX/TensorRT 部署'
)

doc.add_paragraph('')
p = doc.add_paragraph()
p.add_run('这个项目的核心特色是 ').font.size = Pt(11)
p.add_run('ultralytics/nn/backbone/ 和 ultralytics/nn/extra_modules/').bold = True
p.add_run(' 中整合的 ').font.size = Pt(11)
p.add_run('上百个前沿创新模块').bold = True
p.add_run('，你只需修改 RT-DETR 的 YAML 配置文件即可更换 Backbone、Neck 或加入各种注意力机制，非常适合做消融实验和论文对比实验。').font.size = Pt(11)

# ============ 调整列宽 ============
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.space_before = Pt(2)
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    if len(table.columns) == 2:
        for row in table.rows:
            row.cells[0].width = Cm(4)
            row.cells[1].width = Cm(12)

# ============ 保存 ============
save_path = 'D:/RT-DETR项目文件结构详解.docx'
doc.save(save_path)
print(f'文档已保存至: {save_path}')
